#!/usr/bin/env python3
"""
Check Word template for placeholders
"""
from docx import Document

def check_template_placeholders():
    """Check which placeholders exist in the Word template"""
    template_path = "Inputs/FT-010334_Report_RPT_Template_copy.docx"
    
    try:
        doc = Document(template_path)
        
        # Check all paragraphs in the document
        all_text = []
        for paragraph in doc.paragraphs:
            all_text.append(paragraph.text)
        
        # Check headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                all_text.append(header.text)
            for footer in section.footer.paragraphs:
                all_text.append(footer.text)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        
        full_text = " ".join(all_text)
        
        # List of placeholders to check
        placeholders = [
            '[BK_TITLE]',
            '[BK_PURPOSE_TEXT]',
            '[BK_SCOPE_TEXT]',
            '[BK_REFERENCES]',
            '[BK_ACRONYMS]',
            '[BK_DEFINITIONS]',
            '[BK_PROCEDURE_SUMMARY]',
            '[BK_DUT_CONFIG]',
            '[BK_TEST_EXECUTION_CHRONOLOGY]',
            '[BK_DUT_TEST_ARTICLE_CONFIG]',
            '[BK_DUT_TEST_ARTICLE_TRACEABILITY_INFO]',
            '[BK_CONSUMABLES_USED]',
            '[BK_TEST_RESULT_SUMMARY]',
            '[BK_TEST_RESULT_ANALYSIS]',
            '[BK_TEST_METHOD_LOSS_INVESTIGATIONS]',
            '[BK_PROTOCOL_DEVIATIONS]',
            '[BK_DEFECTIVE_UNIT]',
            '[BK_CONCLUSION]'
        ]
        
        print("=== WORD TEMPLATE PLACEHOLDER CHECK ===")
        print(f"Template: {template_path}")
        print()
        
        found_placeholders = []
        missing_placeholders = []
        
        for placeholder in placeholders:
            if placeholder in full_text:
                found_placeholders.append(placeholder)
                print(f"✅ FOUND: {placeholder}")
            else:
                missing_placeholders.append(placeholder)
                print(f"❌ MISSING: {placeholder}")
        
        print()
        print(f"Summary: {len(found_placeholders)} found, {len(missing_placeholders)} missing")
        
        if missing_placeholders:
            print("\nMissing placeholders:")
            for placeholder in missing_placeholders:
                print(f"  - {placeholder}")
        
        return found_placeholders, missing_placeholders
        
    except Exception as e:
        print(f"Error checking template: {str(e)}")
        return [], []

if __name__ == "__main__":
    check_template_placeholders()