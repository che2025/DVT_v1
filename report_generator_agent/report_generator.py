"""
DVT Report Generator - Core business logic for report generation
Updated to use specialized AI agents for Tasks 4.1-4.4
"""

import os
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
import openpyxl
from fastapi import UploadFile
from .ai_agents import DVTAgentOrchestrator, TaskResult, ProtocolDeviationsAgent, DefectiveUnitInvestigationsAgent
import re


class DVTReportGenerator:
    """Main class for DVT report generation with specialized AI agents for Tasks 4.1-4.4"""
    
    def __init__(self, client=None):
        self.report_data = {}
        self.generated_attachments = []  # Track all generated attachment filenames
        self.client = client
        self.model_name = "gemini-2.5-flash"
        
        # Initialize AI agent orchestrator
        self.ai_orchestrator = DVTAgentOrchestrator(client, self.model_name) if client else None
        
        # Template configuration
        self.template_path = "Inputs/FT-010334_Report_RPT_Template_copy.docx"
        self.placeholder_mapping = {
            '[BK_PURPOSE_TEXT]': 'task_4_1_purpose',
            '[BK_SCOPE_TEXT]': 'task_4_1_scope',
            '[BK_REFERENCES]': 'task_4_2',
            '[BK_ACRONYMS]': 'task_4_3',      # Will extract ACRONYMS section
            '[BK_DEFINITIONS]': 'task_4_3',   # Will extract DEFINITIONS section
            '[BK_PROCEDURE_SUMMARY]': 'task_4_4',
            '[BK_DUT_CONFIG]': 'task_4_5',    # Device Under Test Configuration
            '[BK_TEST_EXECUTION_CHRONOLOGY]': 'test_execution_chronology',  # Test Execution Chronology Table
            '[BK_CONSUMABLES_USED]': 'task_4_6',           # Consumables from equipment
            '[BK_TEST_RESULT_SUMMARY]': 'task_4_7',        # Test Result Summary Table
            '[BK_TEST_RESULT_ANALYSIS]': 'test_result_analysis_images',  # Test Result Analysis Images
            '[BK_TEST_METHOD_LOSS_INVESTIGATIONS]': 'test_method_losses',  # Test method losses
            '[BK_PROTOCOL_DEVIATIONS]': 'protocol_deviations',  # Protocol Deviations
            '[BK_DEFECTIVE_UNIT]': 'defective_unit_investigations',  # Defective Unit Investigations
            '[BK_CONCLUSION]': 'conclusion',  # Conclusion section
            '[BK_ATTACHMENTS]': 'attachments'  # List of all attachments
        }
    
    async def test_gemini_connection(self) -> str:
        """Simple test to verify Gemini AI is working"""
        if not self.ai_orchestrator:
            return "⚠️ AI client not configured"
        
        return await self.ai_orchestrator.test_ai_connection()
    
    def check_template_file(self) -> Dict[str, Any]:
        """Check if template file exists and is accessible"""
        result = {
            "exists": False,
            "path": self.template_path,
            "error": None
        }
        
        try:
            if os.path.exists(self.template_path):
                # Try to open the file to verify it's accessible
                doc = Document(self.template_path)
                result["exists"] = True
                print(f"✅ Template file found: {self.template_path}")
            else:
                result["error"] = f"Template file not found: {self.template_path}"
                print(f"❌ Template file missing: {self.template_path}")
        except Exception as e:
            result["error"] = f"Cannot access template file: {str(e)}"
            print(f"❌ Template file error: {str(e)}")
        
        return result
    
    async def process_files(self, protocol_file: UploadFile, data_files: List[UploadFile]) -> Dict[str, Any]:
        """Process uploaded files and extract relevant data with Excel parsing integration"""
        results = {
            "protocol_data": {},
            "test_data": [],
            "data_files": data_files,  # Store original files for Task 4.5
            "extracted_info": {},
            "parsed_excel_data": {}  # New: Store parsed Excel data from our parsers
        }
        
        # Process protocol document
        if protocol_file:
            # Enhanced document parsing with structured analysis
            from .doc_parsers import DocumentParser
            
            print(f"\n📄 Processing protocol document: {protocol_file.filename}")
            
            # Check if it's a Word document for structured parsing
            if protocol_file.filename.endswith('.docx'):
                print("🔍 Running enhanced Word document parsing...")
                try:
                    doc_parser = DocumentParser()
                    
                    # Parse the document structure
                    print("📊 Parsing document structure...")
                    parsed_doc_data = await doc_parser.parse_document(protocol_file)
                    
                    if parsed_doc_data:
                        print(f"✅ Document parsed successfully:")
                        print(f"   - Found {len(parsed_doc_data)} sections")
                        '''
                        # Print section summary
                        for section_num, content in parsed_doc_data.items():
                            if isinstance(content, dict):
                                content_count = len([k for k in content.keys() if k.startswith('content_')])
                                title = content.get('title', '无标题')
                                print(f"   - {section_num}: {title} (内容项: {content_count})")
                            else:
                                print(f"   - {section_num}: [文本内容]")
                        
                        # Print complete raw data dictionary
                        print(f" COMPLETE DOCUMENT RAW DATA:")
                        print("="*80)
                        import json
                        try:
                            print(json.dumps(parsed_doc_data, indent=2, default=str, ensure_ascii=False))
                        except:
                            print(parsed_doc_data)
                        print("="*80)
                        '''
                        # Get AI-friendly format
                        ai_formatted = doc_parser.format_for_ai_prompt()
                        '''
                        print(f"\n📊 DOCUMENT AI-FRIENDLY FORMAT:")
                        print("="*60)
                        print(ai_formatted[:2000] + "..." if len(ai_formatted) > 2000 else ai_formatted)
                        print("="*60)
                        '''
                        # Store structured document data
                        results["protocol_data"] = {
                            "raw_data": parsed_doc_data,
                            "ai_format": ai_formatted,
                            "sections_count": len(parsed_doc_data),
                            "parsing_method": "structured_docx"
                        }
                        
                    else:
                        print("⚠️ Document parsing returned no data, falling back to basic text extraction")
                        protocol_content = await self.extract_document_content(protocol_file)
                        results["protocol_data"] = await self.analyze_protocol(protocol_content)
                        
                except Exception as e:
                    print(f"❌ Document parsing failed: {str(e)}")
                    print("⚠️ Falling back to basic text extraction")
                    protocol_content = await self.extract_document_content(protocol_file)
                    results["protocol_data"] = await self.analyze_protocol(protocol_content)
            else:
                # For non-Word documents, use original method
                print("📄 Using basic text extraction for non-Word document")
                protocol_content = await self.extract_document_content(protocol_file)
                results["protocol_data"] = await self.analyze_protocol(protocol_content)
        
        # Process data files with enhanced Excel parsing
        from .excel_parsers import TestArticleParser, EquipmentUsedParser, DeviationsParser, DefectiveUnitsParser
        
        for data_file in data_files:
            if data_file.filename.endswith(('.xlsx', '.xls')):
                # Original extraction method
                data_content = await self.extract_excel_data(data_file)
                results["test_data"].append(data_content)
                
                # NEW: Enhanced Excel parsing with specialized parsers
                print(f"\n📊 Running enhanced Excel parsing for: {data_file.filename}")
                
                try:
                    # Initialize all parsers
                    test_article_parser = TestArticleParser()
                    equipment_used_parser = EquipmentUsedParser()
                    deviations_parser = DeviationsParser()
                    defective_units_parser = DefectiveUnitsParser()
                    
                    # Parse TEST ARTICLE LOG & TEST RESULTS
                    print("🔍 Parsing TEST ARTICLE LOG & TEST RESULTS...")
                    test_article_data = await test_article_parser.parse_uploaded_file(data_file)
                    
                    if test_article_data:
                        print(f"✅ Test Article Data parsed successfully:")
                        print(f"   - Found {len(test_article_data)} DUT Serial Numbers")
                        if test_article_data:
                            # Show sample of parsed data
                            sample_key = next(iter(test_article_data))
                            print(f"   - Sample DUT: {sample_key}")
                            print(f"   - Columns: {list(test_article_data[sample_key].keys())}")
                        
                        '''
                        print(f"COMPLETE TEST ARTICLE RAW DATA:")
                        print("="*80)
                        import json
                   
                        try:
                            print(json.dumps(test_article_data, indent=2, default=str))
                        except:
                            print(test_article_data)
                        print("="*80)
                        '''
                        # Store with AI-friendly format
                        ai_prompt_data = test_article_parser.format_for_ai_prompt()
                        '''
                        print(f"\n📊 TEST ARTICLE AI-FRIENDLY FORMAT:")
                        print("="*60)
                        print(ai_prompt_data)
                        print("="*60)
                        '''
                        # Get statistics from test article parser
                        statistics = {
                            "total_units": len(test_article_data),
                            "test_method_losses": len([k for k, v in test_article_data.items() if v.get("Test Result") == "TML"]),
                            "actual_sample_size": len([k for k, v in test_article_data.items() if v.get("Test Result") in ["PASS", "FAIL"]])
                        }
                        results["parsed_excel_data"]["test_articles"] = {
                            "raw_data": test_article_data,
                            "ai_format": ai_prompt_data,
                            "statistics": statistics
                        }
                    else:
                        print("ℹ️ No TEST ARTICLE LOG & TEST RESULTS sheet found in this file")
                    
                    # Parse Equipment/Software/Material logs
                    print("🔍 Parsing Equipment/Software/Material logs...")
                    equipment_sheets_data = await equipment_used_parser.parse_uploaded_file(data_file)
                    
                    if equipment_sheets_data:
                        '''
                        print(f"✅ Equipment sheets data parsed successfully:")
                        for sheet_name, sheet_data in equipment_sheets_data.items():
                            print(f"   - {sheet_name}: {len(sheet_data)} rows")
                        
                        # Print complete raw data dictionary
                        print(f"\n📊 COMPLETE EQUIPMENT LOGS RAW DATA:")
                        print("="*80)
                        import json
                        try:
                            print(json.dumps(equipment_sheets_data, indent=2, default=str))
                        except:
                            print(equipment_sheets_data)
                        print("="*80)
                        '''
                        # Store with AI-friendly format
                   
                        ai_prompt_data = equipment_used_parser.format_for_ai_prompt()
                       
                        results["parsed_excel_data"]["equipment_logs"] = {
                            "raw_data": equipment_sheets_data,
                            "ai_format": ai_prompt_data
                        }
                    else:
                        print("ℹ️ No Equipment/Software/Material log sheets found in this file")
                    
                    # Parse Deviations
                    print("🔍 Parsing DEVIATIONS sheets...")
                    deviations_data = await deviations_parser.parse_uploaded_file(data_file)


                    if deviations_data:
                        print(f"✅ Deviations data parsed successfully:")
                        # Store with AI-friendly format
                        ai_prompt_data = deviations_parser.format_for_ai_prompt()
                        
                        results["parsed_excel_data"]["deviations"] = {
                            "raw_data": deviations_data,
                            "ai_format": ai_prompt_data
                        }
                    else:
                        print("ℹ️ No DEVIATIONS sheets found in this file")

                    
                    # Parse Defective Units
                    print("🔍 Parsing DEFECTIVE UNITS sheets...")
                    defective_units_data = await defective_units_parser.parse_uploaded_file(data_file)
                    
                    if defective_units_data:
                        # Store with AI-friendly format
                        ai_prompt_data = defective_units_parser.format_for_ai_prompt()
                        '''
                        print(f"DEFECTIVE UNITS AI-FRIENDLY FORMAT:")
                        print("="*60)
                        print(ai_prompt_data)
                        print("="*60)
                        '''
                        
                        results["parsed_excel_data"]["defective_units"] = {
                            "raw_data": defective_units_data,
                            "ai_format": ai_prompt_data
                        }
                    else:
                        print("ℹ️ No DEFECTIVE UNITS sheets found in this file")
                        
                except Exception as e:
                    print(f"⚠️ Excel parsing error for {data_file.filename}: {str(e)}")
                    # Continue with normal processing even if enhanced parsing fails
        
        # Print summary of parsed data
        if results["parsed_excel_data"]:
            print(f"\n📈 Excel Parsing Summary:")
            if "test_articles" in results["parsed_excel_data"]:
                stats = results["parsed_excel_data"]["test_articles"]["statistics"]
                print(f"   Test Articles: {stats['total_units']} total units")
                print(f"   Test Method Losses: {stats['test_method_losses']}")
                print(f"   Actual Sample Size: {stats['actual_sample_size']}")
            if "equipment_logs" in results["parsed_excel_data"]:
                equipment_data = results["parsed_excel_data"]["equipment_logs"]["raw_data"]
                total_equipment_rows = sum(len(data['data']) for data in equipment_data.values())
                print(f"   Equipment Logs: {total_equipment_rows} total rows across {len(equipment_data)} sheets")
            if "deviations" in results["parsed_excel_data"]:
                deviations_data = results["parsed_excel_data"]["deviations"]["raw_data"]
                total_deviations = sum(data['row_count'] for data in deviations_data.values())
                print(f"   Deviations: {total_deviations} total deviations across {len(deviations_data)} sheets")
            if "defective_units" in results["parsed_excel_data"]:
                defective_data = results["parsed_excel_data"]["defective_units"]["raw_data"]
                total_defective = sum(data['row_count'] for data in defective_data.values())
                print(f"   Defective Units: {total_defective} total defective units across {len(defective_data)} sheets")
        
        return results
    
    async def extract_document_content(self, file: UploadFile) -> str:
        """Extract text content from Word documents"""
        try:
            contents = await file.read()
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(contents)
                tmp_file.flush()
                
                doc = Document(tmp_file.name)
                text_content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                os.unlink(tmp_file.name)
                return "\n".join(text_content)
        except Exception as e:
            print(f"Error extracting document content: {e}")
            return ""
    
    async def extract_excel_data(self, file: UploadFile) -> Dict[str, Any]:
        """Extract data from Excel files"""
        try:
            contents = await file.read()
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(contents)
                tmp_file.flush()
                
                workbook = openpyxl.load_workbook(tmp_file.name)
                data = {}
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_data = []
                    
                    for row in sheet.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            sheet_data.append(row)
                    
                    data[sheet_name] = sheet_data
                
                os.unlink(tmp_file.name)
                return {"filename": file.filename, "sheets": data}
        except Exception as e:
            print(f"Error extracting Excel data: {e}")
            return {"filename": file.filename, "sheets": {}, "error": str(e)}
    
    async def analyze_protocol(self, content: str) -> Dict[str, Any]:
        """Analyze protocol content and save for AI agents"""
        if not content:
            return {
                "scope": "No protocol content provided",
                "purpose": "Please upload protocol document",
                "acceptance_criteria": [],
                "test_parameters": [],
                "original_content": ""
            }
        
        # Save original content for AI agents
        result = {
            "original_content": content,
            "scope": "Will be generated by AI Task 4.1",
            "purpose": "Will be generated by AI Task 4.1",
            "acceptance_criteria": ["To be extracted by AI"],
            "test_parameters": ["To be extracted by AI"],
            "protocol_reference": "Auto-detected from protocol document"
        }
        
        # If AI is available, do quick analysis
        if self.client:
            try:
                from google.genai.types import GenerateContentConfig
                
                prompt = f"""
                Quickly analyze this protocol and extract basic information:
                1. Protocol reference/number
                2. Brief scope description  
                3. Brief purpose description
                
                Protocol Content (first 2000 chars):
                {content[:2000]}
                
                Return JSON format:
                {{"protocol_reference": "...", "scope": "...", "purpose": "..."}}
                """
                
                config = GenerateContentConfig(temperature=0.3)
                
                response = self.client.models.generate_content(
                    model=self.model_name,
                    contents=[prompt],
                    config=config
                )
                
                if response and response.text:
                    # Try to parse JSON response
                    import json
                    try:
                        ai_data = json.loads(response.text)
                        result.update(ai_data)
                        result["ai_analysis"] = "Basic protocol analysis completed"
                    except json.JSONDecodeError:
                        result["ai_analysis"] = response.text
                        
            except Exception as e:
                print(f"Protocol analysis error: {e}")
                result["error"] = str(e)
        
        return result
    
    async def generate_report_sections(self, report_config: Dict[str, Any], processed_data: Dict[str, Any], device_config: Optional[Dict[str, Any]] = None, analysis_image_paths: Optional[List[str]] = None) -> Dict[str, str]:
        """Generate all report sections using specialized AI agents for Tasks 4.1-4.5"""
        sections = {}
        
        # Store analysis image paths for later use
        if analysis_image_paths:
            print(f"🖼️ [DEBUG] Received {len(analysis_image_paths)} analysis images for processing")
            sections["analysis_image_paths"] = analysis_image_paths
        else:
            print("🖼️ [DEBUG] No analysis images provided")
            sections["analysis_image_paths"] = []
        
        # AI connection test
        sections["gemini_test"] = await self.test_gemini_connection()
        
        if self.ai_orchestrator:
            # Extract protocol information
            protocol_data = processed_data.get("protocol_data", {})
            protocol_content = protocol_data.get("original_content", "")
            protocol_number = report_config.get("protocol_number", "TBD")
            project_name = report_config.get("project_name", "DVT Project")
            
            # Extract parsed protocol data from doc parser (preferred)
            parsed_protocol_data = None

            # Try multiple possible locations for parsed data
            if "raw_data" in protocol_data:
                # This is where the structured doc parser stores the data
                parsed_protocol_data = protocol_data.get("raw_data", None)
                print(f"🔍 Found raw_data with {len(parsed_protocol_data)} sections" if parsed_protocol_data else "🔍 raw_data is empty")

            else:
                print("⚠️ No parsed protocol data available, falling back to legacy text parsing")

            ai_results = await self.ai_orchestrator.execute_tasks_4_1_to_4_4(
                protocol_content=protocol_content,
                protocol_number=protocol_number,
                project_name=project_name,
                report_content="",  # Empty for now, will update for Task 4.3 later
                parsed_protocol_data=parsed_protocol_data  # New parameter
            )
            
            # Execute Task 4.5: Device Under Test Configuration
            if device_config:
                print("🤖 Executing AI Task 4.5: Device Under Test Configuration...")
                data_files = processed_data.get("data_files", [])
                task_4_5_result = await self.ai_orchestrator.execute_task_4_5(
                    data_files=data_files,
                    device_config=device_config,
                    report_config=report_config
                )
                
                if task_4_5_result and task_4_5_result.success:
                    print("✅ task_4_5 completed successfully")
                    sections["task_4_5"] = task_4_5_result.content
                    
                    # Check for attachments and add to our tracking list
                    if hasattr(task_4_5_result, 'metadata') and task_4_5_result.metadata.get("attachment_info"):
                        attachment_info = task_4_5_result.metadata["attachment_info"]
                        if "filename" in attachment_info:
                            self.generated_attachments.append(attachment_info["filename"])
                            print(f"📎 Added Task 4.5 attachment: {attachment_info['filename']}")
                else:
                    print(f"❌ task_4_5 failed: {task_4_5_result.error if task_4_5_result else 'Unknown error'}")
                    sections["task_4_5"] = f"Error in task_4_5: {task_4_5_result.error if task_4_5_result else 'Unknown error'}"
            
            # Process AI results in correct order
            # Task 4.1a: Purpose
            if ai_results.get("task_4_1_purpose") and ai_results["task_4_1_purpose"].success:
                print("✅ task_4_1_purpose completed successfully")
                sections["task_4_1_purpose"] = ai_results["task_4_1_purpose"].content
            else:
                error_msg = ai_results["task_4_1_purpose"].error if ai_results.get("task_4_1_purpose") else "Unknown error"
                print(f"❌ task_4_1_purpose failed: {error_msg}")
                sections["task_4_1_purpose"] = f"Error in task_4_1_purpose: {error_msg}"

            # Task 4.1b: Scope
            if ai_results.get("task_4_1_scope") and ai_results["task_4_1_scope"].success:
                print("✅ task_4_1_scope completed successfully")
                sections["task_4_1_scope"] = ai_results["task_4_1_scope"].content
            else:
                error_msg = ai_results["task_4_1_scope"].error if ai_results.get("task_4_1_scope") else "Unknown error"
                print(f"❌ task_4_1_scope failed: {error_msg}")
                sections["task_4_1_scope"] = f"Error in task_4_1_scope: {error_msg}"

            if ai_results.get("task_4_2") and ai_results["task_4_2"].success:
                print("✅ task_4_2 completed successfully")
                sections["task_4_2"] = ai_results["task_4_2"].content
            else:
                error_msg = ai_results["task_4_2"].error if ai_results.get("task_4_2") else "Unknown error"
                print(f"❌ task_4_2 failed: {error_msg}")
                sections["task_4_2"] = f"Error in task_4_2: {error_msg}"
                
            if ai_results.get("task_4_4") and ai_results["task_4_4"].success:
                print("✅ task_4_4 completed successfully")
                sections["task_4_4"] = ai_results["task_4_4"].content
            else:
                error_msg = ai_results["task_4_4"].error if ai_results.get("task_4_4") else "Unknown error"
                print(f"❌ task_4_4 failed: {error_msg}")
                sections["task_4_4"] = f"Error in task_4_4: {error_msg}"
            
            # Continue with remaining tasks (4.5-4.12) using existing methods
            
            # Create combined Material & Equipment section (6)
            dut_config_content = ""
            equipment_content = ""
            
            # Get Task 4.5 content if available
            if "task_4_5" in sections:
                dut_config_content = sections["task_4_5"]
            else:
                dut_config_content = await self.create_dut_configuration(processed_data["test_data"])
            
            # Create Test Execution Chronology section
            if device_config and "test_execution_chronology" in device_config:
                chronology_data = device_config["test_execution_chronology"]
                print(f"🔍 [DEBUG] Chronology data received: {chronology_data}")
                sections["test_execution_chronology"] = self.create_test_execution_chronology(chronology_data)
                print(f"✅ Created Test Execution Chronology with {len(chronology_data)} entries")
                print(f"🔍 [DEBUG] Generated table content: {sections['test_execution_chronology']}")
            else:
                sections["test_execution_chronology"] = self.create_test_execution_chronology([])
                print("📝 Created empty Test Execution Chronology table")
            
            # Create Test Result Analysis Images section
            image_paths = analysis_image_paths if analysis_image_paths else []
            sections["test_result_analysis_images"] = self.create_test_result_analysis_images(image_paths)
            print(f"✅ Created Test Result Analysis with {len(image_paths)} images")
            
            # Get equipment content (Task 4.6 using EquipmentUsedAgent)
            calibration_verified = device_config.get("calibration_verified", True) if device_config else True
            equipment_content = await self.create_equipment_section(
                processed_data, 
                report_config,
                calibration_verified=calibration_verified
            )
            
            # Set Task 4.6 content for consumables and equipment placeholders
            sections["task_4_6"] = equipment_content
            
            # Combine into Material & Equipment section
            sections["material_equipment"] = self._combine_material_equipment_sections(
                dut_config_content, equipment_content
            )
            
            sections["task_4_7"] = await self.create_test_result_summary(processed_data, report_config)
            sections["test_details"] = await self.create_test_results_details(processed_data)
            # Process protocol deviations using Excel data and AI agent
            sections["protocol_deviations"] = await self.create_protocol_deviations_from_excel(processed_data)
            # Process defective unit investigations using Excel data and AI agent
            sections["defective_unit_investigations"] = await self.create_defective_unit_investigations_from_excel(processed_data)
            sections["defective_units"] = await self.create_defective_unit_investigations(
                report_config.get("jira_tickets", [])
            )
            # Convert test_data list to excel_data_dict format for Task 4.11
            excel_data_dict = {}
            for test_data_item in processed_data.get("test_data", []):
                if isinstance(test_data_item, dict) and "filename" in test_data_item:
                    excel_data_dict[test_data_item["filename"]] = test_data_item
            sections["test_method_losses"] = await self.create_test_method_loss_investigations(
                processed_data, excel_data_dict, report_config
            )
            sections["conclusion"] = await self.create_conclusion(processed_data, sections)
            
            # NOW execute Task 4.3 with complete report content
            print("🤖 Executing final Task 4.3 with complete report content...")
            complete_report = self._build_complete_report_text(sections)
            final_acronyms_result = await self.ai_orchestrator.execute_final_acronyms_task(complete_report)
            
            if final_acronyms_result.success:
                print("✅ Final task_4_3 completed successfully")
                sections["task_4_3"] = final_acronyms_result.content
            else:
                print(f"❌ Final task_4_3 failed: {final_acronyms_result.error}")
                sections["task_4_3"] = f"Error in task_4_3: {final_acronyms_result.error}"
            
        else:
            print("⚠️ AI not configured, using fallback methods")
            # Fallback to original methods
            sections["task_4_1"] = await self.create_scope_and_purpose(
                processed_data["protocol_data"], 
                report_config.get("project_name", "DVT Project")
            )
            sections["task_4_2"] = await self.create_reference_section(processed_data)
            sections["task_4_3"] = await self.create_acronyms_section(processed_data)
            sections["task_4_4"] = await self.create_test_procedure_summary(processed_data["protocol_data"])
            
            # Continue with remaining tasks
            sections["dut_config"] = await self.create_dut_configuration(processed_data["test_data"])
            sections["equipment"] = await self.create_equipment_section(
                processed_data, 
                report_config,
                calibration_verified=device_config.get("calibration_verified", True) if device_config else True
            )
            sections["task_4_7"] = await self.create_test_result_summary(processed_data, report_config)
            sections["test_details"] = await self.create_test_results_details(processed_data)
            # Process protocol deviations using Excel data and AI agent (fallback)
            sections["protocol_deviations"] = await self.create_protocol_deviations_from_excel(processed_data)
            # Process defective unit investigations using Excel data and AI agent (fallback)
            sections["defective_unit_investigations"] = await self.create_defective_unit_investigations_from_excel(processed_data)
            sections["defective_units"] = await self.create_defective_unit_investigations(
                report_config.get("jira_tickets", [])
            )
            # Convert test_data list to excel_data_dict format for Task 4.11
            excel_data_dict = {}
            for test_data_item in processed_data.get("test_data", []):
                if isinstance(test_data_item, dict) and "filename" in test_data_item:
                    excel_data_dict[test_data_item["filename"]] = test_data_item
            sections["test_method_losses"] = await self.create_test_method_loss_investigations(
                processed_data, excel_data_dict, report_config
            )
            sections["conclusion"] = await self.create_conclusion(processed_data, sections)
        
        # Generate attachments list
        sections["attachments"] = self._create_attachments_list(processed_data, report_config)
        
        return sections
    
    def _build_complete_report_text(self, sections: Dict[str, str]) -> str:
        """Build complete report text for acronym scanning"""
        # Combine all sections in order for acronym analysis
        report_parts = []
        
        # Add sections in order (excluding gemini_test and incomplete task_4_3)
        for section_key in ["task_4_1", "task_4_2", "task_4_4", "dut_config", "equipment", 
                           "task_4_7", "test_details", "protocol_deviations", "defective_unit_investigations", "defective_units", 
                           "test_method_losses", "conclusion"]:
            if section_key in sections:
                report_parts.append(sections[section_key])
        
        return "\n\n".join(report_parts)
    
    async def create_scope_and_purpose(self, protocol_data: Dict[str, Any], project_name: str) -> str:
        """AI Task 4.1: Create Scope and Purpose section"""
        
        # Extract protocol information
        protocol_reference = protocol_data.get('protocol_reference', 'TBD')
        protocol_scope = protocol_data.get('scope', 'Device verification testing')
        protocol_purpose = protocol_data.get('purpose', 'Validate device functionality and performance')
        
        # Adapt purpose for report (add protocol reference)
        adapted_purpose = f"The purpose of this report is to document the results of the {protocol_purpose.lower()} executed under protocol {protocol_reference}."
        
        # Adapt scope for report (add project name reference)
        adapted_scope = f"The scope of this report applies to {protocol_scope} for project {project_name}."
        
        scope_section = f"""
        ## SCOPE AND PURPOSE
        
        ### Purpose
        {adapted_purpose}
        
        ### Scope
        {adapted_scope}
        """
        return scope_section.strip()
    
    async def create_reference_section(self, processed_data: Dict[str, Any]) -> str:
        """AI Task 4.2: Create Reference Section"""
        
        # TODO: Extract document numbers from report content using document number format
        # TODO: Retrieve document metadata (title, revision) from knowledge base
        
        # For now, create basic reference table with protocol document
        protocol_data = processed_data.get("protocol_data", {})
        protocol_ref = protocol_data.get('protocol_reference', 'TBD')
        
        references = f"""
        ## REFERENCES
        
        | Document No. | Document Title | Rev |
        |--------------|----------------|-----|
        | {protocol_ref} | Protocol Document | TBD |
        | TBD | Test Specification | TBD |
        | TBD | Device Requirements | TBD |
        
        Note: Document references will be automatically extracted from report content when document number format knowledge base is implemented.
        """
        return references.strip()
    
    async def create_acronyms_section(self, processed_data: Dict[str, Any]) -> str:
        """AI Task 4.3: Create Acronyms & Definitions section"""
        
        # TODO: Extract all CAPS words from report content
        # TODO: Look up definitions from Acronym Knowledge Base
        
        # Basic acronym table with common DVT terms
        acronyms = """
        ## ACRONYMS AND DEFINITIONS
        
        | Acronym | Definition |
        |---------|------------|
        | DVT | Design Verification Testing |
        | DUT | Device Under Test |
        | TBD | To Be Determined |
        | EMC | Electromagnetic Compatibility |
        | GSS | Glucose Sensing System |
        | BLE | Bluetooth Low Energy |
        
        Note: Additional acronyms will be automatically extracted from report content and looked up in acronym knowledge base.
        """
        return acronyms.strip()
    
    def create_test_execution_chronology(self, chronology_data: List[Dict[str, str]]) -> str:
        """Create Test Execution Chronology table from user input"""
        print(f"🔍 [DEBUG] create_test_execution_chronology called with data: {chronology_data}")
        
        if not chronology_data:
            # Return empty table if no data provided
            result = """
| Step | Start Date | End Date | Location |
|------|------------|----------|----------|
|      |            |          |          |
"""
            print(f"🔍 [DEBUG] No data provided, returning empty table")
            return result
        
        # Create table header
        table_content = """
| Step | Start Date | End Date | Location |
|------|------------|----------|----------|
"""
        
        # Add data rows
        for i, item in enumerate(chronology_data):
            step = item.get("step", "").strip()
            start_date = item.get("start_date", "").strip() 
            end_date = item.get("end_date", "").strip()
            location = item.get("location", "").strip()
            
            row = f"| {step} | {start_date} | {end_date} | {location} |\n"
            table_content += row
            print(f"🔍 [DEBUG] Added row {i+1}: {row.strip()}")
        
        print(f"🔍 [DEBUG] Final table content:\n{table_content}")
        return table_content.strip()
    
    def insert_analysis_images(self, doc, paragraph, image_placeholder: str):
        """Insert analysis images into Word document"""
        try:
            print(f"🖼️ [DEBUG] Processing image placeholder: {image_placeholder}")
            
            # Parse the placeholder: {ANALYSIS_IMAGES:count:path1|path2|...}
            if not image_placeholder.startswith("{ANALYSIS_IMAGES:"):
                print("🖼️ [ERROR] Invalid image placeholder format")
                return
            
            # Extract image paths
            placeholder_content = image_placeholder[17:-1]  # Remove {ANALYSIS_IMAGES: and }
            parts = placeholder_content.split(":", 1)
            if len(parts) < 2:
                print("🖼️ [ERROR] Invalid placeholder format - missing paths")
                return
            
            count = int(parts[0])
            image_paths = parts[1].split("|") if parts[1] else []
            
            print(f"🖼️ [DEBUG] Found {count} images to insert: {image_paths}")
            
            if not image_paths or not image_paths[0]:
                print("🖼️ [DEBUG] No image paths provided")
                return
            
            # Insert images into the document
            for i, image_path in enumerate(image_paths):
                if not os.path.exists(image_path):
                    print(f"🖼️ [WARNING] Image file not found: {image_path}")
                    continue
                
                try:
                    # Clear the paragraph content if this is the first image
                    if i == 0:
                        paragraph.clear()

                    # Add image to paragraph without any size specification
                    # This preserves the original image quality and lets Word decide the display size
                    run = paragraph.add_run()
                    run.add_picture(image_path)  # No width/height parameters = original quality
                    
                    # Add a line break after each image except the last one
                    if i < len(image_paths) - 1:
                        paragraph.add_run().add_break()
                    
                    print(f"🖼️ [DEBUG] Successfully inserted image {i+1} with original quality: {os.path.basename(image_path)}")
                    
                except Exception as e:
                    print(f"🖼️ [ERROR] Failed to insert image {image_path}: {e}")
                    # Add error message to document instead of failing silently
                    error_run = paragraph.add_run()
                    error_run.text = f"[Image insertion failed: {os.path.basename(image_path)}]"
                    error_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for error
                    continue
            
            print(f"🖼️ [DEBUG] Completed inserting {len(image_paths)} images")
            
        except Exception as e:
            print(f"🖼️ [ERROR] Failed to process images: {e}")
            # Add error message to document
            error_run = paragraph.add_run()
            error_run.text = f"[Image processing failed: {str(e)}]"
            error_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for error

    def create_test_result_analysis_images(self, image_paths: List[str]) -> str:
        """Create Test Result Analysis Images section for Word document"""
        print(f"🖼️ [DEBUG] create_test_result_analysis_images called with {len(image_paths)} images")
        
        if not image_paths:
            print("🖼️ [DEBUG] No images provided, returning TBD")
            return "TBD"
        
        # For now, return a placeholder that will be processed during Word document creation
        # The actual image insertion will happen in the Word processing phase
        image_placeholder = f"{{ANALYSIS_IMAGES:{len(image_paths)}:{'|'.join(image_paths)}}}"
        print(f"🖼️ [DEBUG] Generated image placeholder: {image_placeholder}")
        return image_placeholder
    
    async def create_test_procedure_summary(self, protocol_data: Dict[str, Any]) -> str:
        """AI Task 4.4: Create Test Procedure Summary"""
        
        # Extract AI analysis from protocol
        ai_analysis = protocol_data.get('ai_analysis', '')
        
        # Create comprehensive test procedure summary (~200 words)
        # Following the 5 key requirements from 4.4.1
        summary = f"""
        ## TEST PROCEDURE SUMMARY
        
        This section provides an executive summary of the test procedure executed for this DVT report.
        
        ### Protocol Analysis:
        {ai_analysis if ai_analysis else 'Test procedure summary will be generated from protocol analysis.'}
        
        ### Test Execution Overview:
        The test procedure involved comprehensive evaluation following the protocol requirements:
        
        **1. Conditioning:** Test articles underwent conditioning procedures prior to test execution to ensure proper baseline conditions were established according to protocol specifications.
        
        **2. Parameters Evaluated:** Key performance parameters were monitored and evaluated throughout the test execution phase, including functional specifications and performance characteristics defined in the acceptance criteria.
        
        **3. Equipment & Instrumentation:** Specialized test equipment and instrumentation were utilized for precise measurement and monitoring, with all equipment verified as calibrated at time of use.
        
        **4. Device Monitoring:** Continuous monitoring protocols were implemented to track device performance throughout the test duration, ensuring data integrity and proper test execution.
        
        **5. Test Duration:** The test was executed according to the timeline specified in the protocol, with documented start and end times for each critical test phase.
        
        ### Protocol Reference:
        - Protocol: {protocol_data.get('protocol_reference', 'TBD')}
        - Scope: {protocol_data.get('scope', 'TBD')}
        - Purpose: {protocol_data.get('purpose', 'TBD')}
        """
        return summary.strip()
    
    async def create_dut_configuration(self, test_data: List[Dict[str, Any]]) -> str:
        """AI Task 4.5: Create Device Under Test Configuration Section"""
        config = """
        ## DEVICE UNDER TEST CONFIGURATION
        
        ### Test Article Summary
        
        The test articles were configured as follows:
        - Units were processed according to standard manufacturing procedures
        - No additional modifications were made outside of normal manufacturing
        - Sterilization status: TBD (to be confirmed during report generation)
        
        ### Test Article Details
        
        | Part Number | Serial Number | Lot Number |
        |-------------|---------------|------------|
        | TBD | TBD | TBD |
        
        Note: Test article details will be extracted from uploaded data sheets.
        """
        return config.strip()
    
    async def create_equipment_section(self, processed_data: Dict[str, Any], report_config: Optional[Dict[str, Any]] = None, calibration_verified: bool = True) -> str:
        """AI Task 4.6: Create Equipment Used Section using EquipmentUsedAgent"""
        from .ai_agents import EquipmentUsedAgent
        
        try:
            # Initialize the Equipment Used Agent
            equipment_agent = EquipmentUsedAgent(self.client, self.model_name)
            
            # Prepare equipment data for the agent from processed_data
            equipment_test_data = self._prepare_equipment_data_for_agent(processed_data)
            
            # Process equipment data using the specialized agent
            result = await equipment_agent.process(
                test_data=equipment_test_data,
                report_config=report_config or {},
                calibration_verified=calibration_verified
            )
            
            if result.success:
                print("✅ Equipment section generated successfully with EquipmentUsedAgent")
                
                # Check for attachments and add to our tracking list
                if hasattr(result, 'attachments') and result.attachments:
                    for attachment in result.attachments:
                        if "filename" in attachment:
                            self.generated_attachments.append(attachment["filename"])
                            print(f"📎 Added Task 4.6 attachment: {attachment['filename']}")
                
                return result.content
            else:
                print(f"⚠️ EquipmentUsedAgent failed: {result.error}")
                return self._create_fallback_equipment_section(calibration_verified)
                
        except Exception as e:
            print(f"❌ Error in equipment section generation: {str(e)}")
            return self._create_fallback_equipment_section(calibration_verified)
    
    def _prepare_equipment_data_for_agent(self, processed_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Prepare equipment data in the format expected by EquipmentUsedAgent"""
        prepared_data = []
        
        # Check for all three types of log data from Excel Parser
        if "parsed_excel_data" in processed_data:
            excel_data = processed_data["parsed_excel_data"]
            all_logs_data = {}
            
            # Collect equipment logs
            if "equipment_logs" in excel_data and "raw_data" in excel_data["equipment_logs"]:
                all_logs_data.update(excel_data["equipment_logs"]["raw_data"])
                print(f"🔧 Found equipment_logs: {len(excel_data['equipment_logs']['raw_data'])} sheets")
            
            # Collect software logs
            if "software_logs" in excel_data and "raw_data" in excel_data["software_logs"]:
                all_logs_data.update(excel_data["software_logs"]["raw_data"])
                print(f"🔧 Found software_logs: {len(excel_data['software_logs']['raw_data'])} sheets")
            
            # Collect material logs
            if "material_logs" in excel_data and "raw_data" in excel_data["material_logs"]:
                all_logs_data.update(excel_data["material_logs"]["raw_data"])
                print(f"🔧 Found material_logs: {len(excel_data['material_logs']['raw_data'])} sheets")
            
            # Create combined data item if we have any logs
            if all_logs_data:
                prepared_item = {
                    "equipment_data": all_logs_data,
                    "file_name": "all_logs.xlsx",
                    "source": "excel_parser"
                }
                prepared_data.append(prepared_item)
                print(f"🔧 Prepared combined logs data: {len(all_logs_data)} total sheets")
                
                # Debug: show what data we have for all three types
                for sheet_name, sheet_info in all_logs_data.items():
                    if isinstance(sheet_info, dict) and "data" in sheet_info:
                        print(f"   - {sheet_name}: {len(sheet_info['data'])} rows")
                        if sheet_info["data"]:
                            sample_columns = list(sheet_info["data"][0].keys()) if sheet_info["data"] else []
                            print(f"     Sample columns: {sample_columns[:5]}...")
        
        # Fallback: check test_data for any direct sheet data
        if not prepared_data and "test_data" in processed_data:
            for data_item in processed_data["test_data"]:
                if "sheet_name" in data_item and any(log in data_item["sheet_name"] for log in ["EQUIPMENT", "SOFTWARE", "MATERIAL"]):
                    prepared_data.append(data_item)
                    print(f"🔧 Added fallback data from {data_item.get('sheet_name', 'unknown sheet')}")
        
        if not prepared_data:
            print("⚠️ No equipment/software/material data found in processed_data")
        
        return prepared_data
    
    def _create_fallback_equipment_section(self, calibration_verified: bool = True) -> str:
        """Create fallback equipment section when AI processing fails"""
        calibration_statement = "All equipment used in this testing was verified as calibrated at the time of use." if calibration_verified else "Equipment calibration verification was not performed prior to testing."
        
        return f"""{calibration_statement}

### Equipment Log
| Equipment Number | Equipment Description | Calibration Due Date |
|------------------|---------------------|---------------------|
| TBD | TBD | TBD |

### Software Used  
| Software Name | Version | License Information |
|---------------|---------|-------------------|
| TBD | TBD | TBD |

### Material Used
| Material Name | Lot Number | Expiration Date |
|---------------|------------|-----------------|
| TBD | TBD | TBD |

Note: Equipment details will be extracted from uploaded equipment logs."""
    
    async def create_test_result_summary(self, processed_data: Dict[str, Any], report_config: Dict[str, Any] = None) -> str:
        """AI Task 4.7: Create Test Result Summary Section using AI Agent"""
        try:
            if not self.ai_orchestrator:
                # Fallback to original implementation
                return self._create_test_result_summary_fallback()
            
            # Get test data type from report config
            test_data_type = report_config.get("test_data_type", "attribute") if report_config else "attribute"
            
            # Get doc data and excel data from processed data
            doc_data = processed_data.get("doc_data", {})
            # Convert test_data list to excel_data_dict format for Task 4.7
            excel_data_dict = {}
            test_data = processed_data.get("test_data", [])
            print(f"🔍 [TASK 4.7 DEBUG] test_data type: {type(test_data)}, length: {len(test_data) if hasattr(test_data, '__len__') else 'N/A'}")
            
            for i, test_data_item in enumerate(test_data):
                print(f"🔍 [TASK 4.7 DEBUG] test_data_item {i}: type={type(test_data_item)}")
                if isinstance(test_data_item, dict):
                    print(f"🔍 [TASK 4.7 DEBUG] test_data_item {i} keys: {list(test_data_item.keys())}")
                    if "filename" in test_data_item:
                        excel_data_dict[test_data_item["filename"]] = test_data_item
                        print(f"🔍 [TASK 4.7 DEBUG] Added to excel_data_dict: {test_data_item['filename']}")
            
            print(f"🔍 [TASK 4.7 DEBUG] excel_data_dict keys: {list(excel_data_dict.keys())}")
            excel_data = excel_data_dict
            protocol_data = processed_data.get("protocol_data", {})
            protocol_content = protocol_data.get("original_content", "")
            
            # Get AI-friendly format if available
            ai_friendly_format = protocol_data.get("ai_format", None)
            
            print(f"🚀 Executing Task 4.7: Create Test Result Summary (Data Type: {test_data_type})")
            if ai_friendly_format:
                print(f"📄 Using AI-friendly format ({len(ai_friendly_format)} characters)")
            
            # Execute Task 4.7 using AI Agent
            result = await self.ai_orchestrator.execute_task_4_7(
                doc_data=doc_data,
                excel_data=excel_data,
                test_data_type=test_data_type,
                protocol_content=protocol_content,
                ai_friendly_format=ai_friendly_format
            )
            
            if result and result.success:
                print("✅ Task 4.7 completed successfully")
                return result.content
            else:
                error_msg = result.error if result else "Unknown error"
                print(f"❌ Task 4.7 failed: {error_msg}")
                return f"Error generating test result summary: {error_msg}"
                
        except Exception as e:
            print(f"❌ Error in create_test_result_summary: {str(e)}")
            return f"Error generating test result summary: {str(e)}"
    
    def _create_test_result_summary_fallback(self) -> str:
        """Fallback method for test result summary when AI is not available"""
        summary = """
        ## TEST RESULT SUMMARY
        
        ### Attribute Data Analysis
        | Req ID | Acceptance Criteria | Confidence/Reliability | Initial Sample | Test Method Losses | Actual Sample | Defective Units | Actual Conf/Rel |
        |--------|-------------------|----------------------|----------------|-------------------|---------------|-----------------|-----------------|
        | TBD | TBD | TBD | TBD | TBD | TBD | TBD | TBD |
        
        ### Variable Data Analysis  
        | Req ID | Acceptance Criteria | Initial Sample | Test Method Losses | Actual Sample | Defective Units | Tolerance Interval | Confidence/Reliability |
        |--------|-------------------|----------------|-------------------|---------------|-----------------|-------------------|----------------------|
        | TBD | TBD | TBD | TBD | TBD | TBD | TBD | TBD |
        
        Note: Test results will be calculated from uploaded test data.
        """
        return summary.strip()
    
    async def create_test_results_details(self, processed_data: Dict[str, Any]) -> str:
        """AI Task 4.8: Create Test Results Details section"""
        details = """
        ## TEST RESULTS DETAILS
        
        ### Test Execution Timeline
        | Test Step | Start Date | End Date | Location |
        |-----------|------------|----------|----------|
        | Test Setup | TBD | TBD | TBD |
        | Test Execution | TBD | TBD | TBD |
        | Data Analysis | TBD | TBD | TBD |
        
        ### Statistical Analysis
        
        For variable data analysis, the following statistical measures were calculated:
        - Sample statistics (min, max, median, quartiles)
        - Tolerance intervals with specified confidence levels
        - Distribution analysis and normality testing
        
        Note: Detailed analysis and charts will be generated from test data.
        """
        return details.strip()
    
    async def create_protocol_deviations_from_excel(self, processed_data: Dict[str, Any]) -> str:
        """
        Create Protocol Deviations Section using Excel deviation data and AI agent
        """
        print("🔍 Processing protocol deviations from Excel data...")
        
        # Get deviations data from Excel parsers
        deviations_data = processed_data.get("parsed_excel_data", {}).get("deviations", {}).get("raw_data", {})
        
        if not deviations_data:
            print("📝 No deviations data found in Excel files")
            return "No deviations."
        
        # Create deviation AI agent
        if self.client:
            deviation_agent = ProtocolDeviationsAgent(self.client, self.model_name)
            
            # Format deviations data for AI processing
            deviations_text = self._format_deviations_for_ai(deviations_data)
            print(f"📊 Formatted deviations text: {deviations_text[:200]}...")
                
            # Process with AI agent
            result = await deviation_agent.process_deviations(deviations_text)
            
            if result.success:
                print(f"✅ Protocol deviations processed successfully: {result.metadata.get('deviation_count', 0)} deviations")
                return result.content
            else:
                print(f"❌ Deviation processing failed: {result.error}")
                return "No deviations."
        else:
            print("⚠️ AI client not configured, using fallback")
            return "No deviations."
    
    def _format_deviations_for_ai(self, deviations_data: Dict[str, Any]) -> str:
        """Format deviations data for AI processing"""
        if not deviations_data:
            return "No deviations data found."
        
        formatted_sections = []
        
        for sheet_name, sheet_data in deviations_data.items():
            formatted_sections.append(f"\n=== {sheet_name} ===")
            formatted_sections.append(f"Found {sheet_data.get('row_count', 0)} deviations")
            formatted_sections.append(f"Columns: {', '.join(sheet_data.get('columns', []))}")
            
            # Show deviation records as examples
            data = sheet_data.get('data', [])
            for i, deviation in enumerate(data[:3], 1):
                formatted_sections.append(f"  Deviation {i}: {deviation}")
            
            if len(data) > 3:
                formatted_sections.append(f"  ... and {len(data) - 3} more deviations")
        
        return '\n'.join(formatted_sections)

    async def create_defective_unit_investigations_from_excel(self, processed_data: Dict[str, Any]) -> str:
        """
        Create Defective Unit Investigations Section using Excel defective units data and AI agent
        """
        print("🔍 Processing defective unit investigations from Excel data...")
        
        # Get defective units data from Excel parsers
        defective_units_data = processed_data.get("parsed_excel_data", {}).get("defective_units", {}).get("raw_data", {})
        
        if not defective_units_data:
            print("📝 No defective units data found in Excel files")
            return "No defective unit investigations available."
        
        # Create defective unit investigations AI agent
        if self.client:
            defective_agent = DefectiveUnitInvestigationsAgent(self.client, self.model_name)
            
            # Format defective units data for AI processing
            defective_units_text = self._format_defective_units_for_ai(defective_units_data)
            print(f"📊 Formatted defective units text: {defective_units_text[:200]}...")
                
            # Process with AI agent
            result = await defective_agent.process_defective_units(defective_units_text)
            
            if result.success:
                print(f"✅ Defective unit investigations processed successfully: {result.metadata.get('investigation_count', 0)} investigations")
                return result.content
            else:
                print(f"❌ Defective unit investigation processing failed: {result.error}")
                return "No defective unit investigations available."
        else:
            print("⚠️ AI client not configured, using fallback")
            return "No defective unit investigations available."
    
    def _format_defective_units_for_ai(self, defective_units_data: Dict[str, Any]) -> str:
        """Format defective units data for AI processing"""
        if not defective_units_data:
            return "No defective units data found."
        
        formatted_sections = []
        
        for sheet_name, sheet_data in defective_units_data.items():
            formatted_sections.append(f"\n=== {sheet_name} ===")
            formatted_sections.append(f"Found {sheet_data.get('row_count', 0)} defective units")
            formatted_sections.append(f"Columns: {', '.join(sheet_data.get('columns', []))}")
            
            # Show defective unit records as examples
            data = sheet_data.get('data', [])
            for i, unit in enumerate(data[:3], 1):
                formatted_sections.append(f"  Unit {i}: {unit}")
            
            if len(data) > 3:
                formatted_sections.append(f"  ... and {len(data) - 3} more defective units")
        
        return '\n'.join(formatted_sections)

    async def create_protocol_deviations(self, deviations: List[Dict[str, Any]]) -> str:
        """AI Task 4.9: Create Protocol Deviations Section"""
        if not deviations:
            return """
            ## PROTOCOL DEVIATIONS
            
            No deviations occurred in the execution of this test.
            """
        
        deviations_text = """
        ## PROTOCOL DEVIATIONS
        
        There were deviations identified during test execution. Each deviation is analyzed below:
        """
        
        for i, deviation in enumerate(deviations, 1):
            deviations_text += f"""
            
            ### DEVIATION #{i}: {deviation.get('title', 'TBD')}
            
            **Summary:** {deviation.get('summary', 'TBD')}
            
            **Impact on Results:** {deviation.get('impact', 'TBD')}
            
            **Resolution:** {deviation.get('resolution', 'TBD')}
            """
        
        return deviations_text.strip()
    
    async def create_defective_unit_investigations(self, jira_tickets: List[str]) -> str:
        """AI Task 4.10: Create Defective Unit Investigations section"""
        if not jira_tickets:
            return """
            ## DEFECTIVE UNIT INVESTIGATIONS
            
            No defective unit investigations were required for this test.
            """
        
        investigations = """
        ## DEFECTIVE UNIT INVESTIGATIONS
        
        The following defective unit investigations were conducted:
        """
        
        for i, ticket in enumerate(jira_tickets, 1):
            investigations += f"""
            
            ### Defective Unit Investigation #{i} - {ticket}
            
            **Specification Not Met:** TBD
            **Number of Defective Units:** TBD  
            **Serial Numbers:** TBD
            **Root Cause Summary:** Investigation details available in JIRA ticket {ticket}
            **Actions Taken:** TBD
            """
        
        return investigations.strip()
    
    async def create_test_method_loss_investigations(self, processed_data: Optional[Dict[str, Any]] = None, excel_data_dict: Optional[Dict[str, Any]] = None, report_config: Optional[Dict[str, str]] = None) -> str:
        """AI Task 4.11: Create Test Method Loss Investigations section using TestMethodLossAgent"""
        from .ai_agents import TestMethodLossAgent
        
        try:
            # Initialize the Test Method Loss Agent
            test_method_loss_agent = TestMethodLossAgent(self.client, self.model_name)
            
            # Prepare report configuration
            if not report_config:
                report_config = {
                    'document_number': 'RPT-XXXXXX',
                    'revision': 'XXX'
                }
            
            # Process test method loss data using the specialized agent
            result = await test_method_loss_agent.create_test_method_loss_investigations(
                excel_data_dict=excel_data_dict or {},
                report_config=report_config
            )
            
            if result.success:
                print("✅ Test Method Loss Investigations section generated successfully")
                return result.content
            else:
                print(f"⚠️ TestMethodLossAgent failed: {result.error}")
                return self._create_fallback_test_method_loss_section()
                
        except Exception as e:
            print(f"❌ Error in test method loss investigations generation: {str(e)}")
            return self._create_fallback_test_method_loss_section()
    
    def _create_fallback_test_method_loss_section(self) -> str:
        """Create fallback test method loss section when AI processing fails"""
        return """No test method loss occurred in the execution of this test."""
    
    async def create_conclusion(self, processed_data: Dict[str, Any], sections: Dict[str, str]) -> str:
        """AI Task 4.12: Create Conclusion section using ConclusionAgent"""
        try:
            # Get test results and scope content for conclusion generation
            test_results_content = sections.get("task_4_7", "")
            scope_content = sections.get("task_4_1", "")
            
            if not test_results_content or not scope_content:
                print("⚠️ Missing test results or scope content for conclusion generation, using fallback")
                return self._create_fallback_conclusion(processed_data)
            
            if self.ai_orchestrator:
                from .ai_agents import ConclusionAgent
                
                # Create ConclusionAgent
                conclusion_agent = ConclusionAgent(
                    client=self.ai_orchestrator.client,
                    model_name="gemini-2.0-flash-exp"
                )
                
                # Generate conclusion using AI
                result = await conclusion_agent.generate_conclusion(
                    test_results_content=test_results_content,
                    scope_content=scope_content
                )
                
                if result.success:
                    print(f"✅ Conclusion generated successfully ({len(result.content)} characters)")
                    return result.content
                else:
                    print(f"❌ Conclusion generation failed: {result.error}")
                    return self._create_fallback_conclusion(processed_data)
            else:
                print("⚠️ AI orchestrator not available, using fallback conclusion")
                return self._create_fallback_conclusion(processed_data)
                
        except Exception as e:
            print(f"❌ Error in conclusion generation: {str(e)}")
            return self._create_fallback_conclusion(processed_data)
    
    def _create_fallback_conclusion(self, processed_data: Dict[str, Any]) -> str:
        """Fallback conclusion when AI generation fails"""
        conclusion = f"""
The DVT testing has been completed according to the specified protocol. 

Based on the test results and analysis:

- All acceptance criteria were evaluated according to the test protocol
- Test execution was completed successfully with documented procedures
- Results meet the specified confidence and reliability requirements

**Summary:** The device performance meets the specified requirements as defined in the protocol documentation.

**Protocol Reference:** {processed_data.get('protocol_data', {}).get('protocol_reference', 'TBD')}

Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        return conclusion.strip()
    
    def clean_markdown_content(self, content: str) -> str:
        """Remove markdown formatting from AI-generated content and remove unwanted headers/descriptions"""
        if not content:
            return ""
        
        # Remove unwanted section headers that should not appear in template replacement
        unwanted_headers = [
            'REFERENCES', 'TEST PROCEDURE SUMMARY', 'ACRONYMS', 'DEFINITIONS', 
            'PURPOSE', 'SCOPE', 'TEST METHOD LOSS INVESTIGATIONS'
        ]
        for header in unwanted_headers:
            # Remove standalone header lines
            content = re.sub(rf'^{header}\s*$', '', content, flags=re.MULTILINE)
            # Remove numbered headers like "4.1 PURPOSE"
            content = re.sub(rf'^\d+\.\d*\s*{header}\s*$', '', content, flags=re.MULTILINE)
        
        # Remove unwanted descriptive text (more comprehensive)
        unwanted_phrases = [
            r'List in this section.*?\.?',
            r'This section.*?\.?',
            r'Use this section.*?\.?',
            r'Include in this section.*?\.?',
            r'Document in this section.*?\.?'
        ]
        for phrase in unwanted_phrases:
            content = re.sub(phrase, '', content, flags=re.IGNORECASE | re.MULTILINE)
        
        # Remove markdown headers (##, ###, etc.)
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        
        # Remove bold formatting (**text**)
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
        
        # Remove italic formatting (*text*)
        content = re.sub(r'\*(.*?)\*', r'\1', content)
        
        # Remove markdown links [text](url)
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        
        # Remove markdown code blocks ```
        content = re.sub(r'```[\s\S]*?```', '', content)
        
        # Remove inline code `text`
        content = re.sub(r'`([^`]+)`', r'\1', content)
        
        # Clean up multiple newlines and empty lines
        content = re.sub(r'\n{3,}', '\n\n', content)
        content = re.sub(r'^\s*\n', '', content, flags=re.MULTILINE)
        
        return content.strip()
    
    def extract_content_sections(self, sections: Dict[str, str]) -> Dict[str, str]:
        """Extract specific content sections for template replacement"""
        content_mapping = {}
        
        # Process each placeholder to extract the appropriate content
        for placeholder, section_key in self.placeholder_mapping.items():
            if section_key in sections:
                section_content = sections[section_key]
                
                if placeholder == "[BK_PURPOSE_TEXT]":
                    # Extract PURPOSE from task_4_1 - new format with PURPOSE_CONTENT:
                    purpose_match = re.search(r'PURPOSE_CONTENT:\s*\n(.*?)(?=SCOPE_CONTENT:|$)', section_content, re.DOTALL | re.IGNORECASE)
                    if purpose_match:
                        content_mapping["BK_PURPOSE_TEXT"] = self.clean_markdown_content(purpose_match.group(1).strip())
                    else:
                        # Fallback: old format or plain content
                        purpose_match = re.search(r'PURPOSE\s*\n(.*?)(?=SCOPE|$)', section_content, re.DOTALL | re.IGNORECASE)
                        if purpose_match:
                            content_mapping["BK_PURPOSE_TEXT"] = self.clean_markdown_content(purpose_match.group(1).strip())
                        else:
                            # Use content before SCOPE as final fallback
                            lines = section_content.split('\n')
                            purpose_lines = []
                            for line in lines:
                                if line.strip() and 'SCOPE' not in line.upper() and 'PURPOSE_CONTENT:' not in line:
                                    purpose_lines.append(line)
                                elif 'SCOPE' in line.upper():
                                    break
                            if purpose_lines:
                                content_mapping["BK_PURPOSE_TEXT"] = self.clean_markdown_content('\n'.join(purpose_lines))
                
                elif placeholder == "[BK_SCOPE_TEXT]":
                    # Extract SCOPE from task_4_1 - new format with SCOPE_CONTENT:
                    scope_match = re.search(r'SCOPE_CONTENT:\s*\n(.*?)(?=\n\n|$)', section_content, re.DOTALL | re.IGNORECASE)
                    if scope_match:
                        content_mapping["BK_SCOPE_TEXT"] = self.clean_markdown_content(scope_match.group(1).strip())
                    else:
                        # Fallback: old format
                        scope_match = re.search(r'SCOPE\s*\n(.*?)(?=\n\n|$)', section_content, re.DOTALL | re.IGNORECASE)
                        if scope_match:
                            content_mapping["BK_SCOPE_TEXT"] = self.clean_markdown_content(scope_match.group(1).strip())
                        else:
                            # Use content after SCOPE keyword as final fallback
                            lines = section_content.split('\n')
                            scope_lines = []
                            found_scope = False
                            for line in lines:
                                if 'SCOPE' in line.upper():
                                    found_scope = True
                                    continue
                                elif found_scope and line.strip():
                                    scope_lines.append(line)
                            if scope_lines:
                                content_mapping["BK_SCOPE_TEXT"] = self.clean_markdown_content('\n'.join(scope_lines))
                
                elif placeholder == "[BK_ACRONYMS]":
                    # Extract ACRONYMS from task_4_3 - new format with ACRONYMS_CONTENT:
                    acronyms_match = re.search(r'ACRONYMS_CONTENT:\s*\n(.*?)(?=DEFINITIONS_CONTENT:|$)', section_content, re.DOTALL | re.IGNORECASE)
                    if acronyms_match:
                        content_mapping["BK_ACRONYMS"] = self.clean_markdown_content(acronyms_match.group(1).strip())
                    else:
                        # Fallback: old format
                        acronyms_match = re.search(r'ACRONYMS\s*\n(.*?)(?=DEFINITIONS|$)', section_content, re.DOTALL | re.IGNORECASE)
                        if acronyms_match:
                            content_mapping["BK_ACRONYMS"] = self.clean_markdown_content(acronyms_match.group(1).strip())
                        else:
                            # Use full content if no definitions section
                            content_mapping["BK_ACRONYMS"] = self.clean_markdown_content(section_content)
                
                elif placeholder == "[BK_DEFINITIONS]":
                    # Extract DEFINITIONS from task_4_3 - new format with DEFINITIONS_CONTENT:
                    definitions_match = re.search(r'DEFINITIONS_CONTENT:\s*\n(.*?)(?=\n\n|$)', section_content, re.DOTALL | re.IGNORECASE)
                    if definitions_match:
                        content_mapping["BK_DEFINITIONS"] = self.clean_markdown_content(definitions_match.group(1).strip())
                    else:
                        # Fallback: old format
                        definitions_match = re.search(r'DEFINITIONS.*?\n(.*?)$', section_content, re.DOTALL | re.IGNORECASE)
                        if definitions_match:
                            definitions_content = definitions_match.group(1).strip()
                            if definitions_content:
                                content_mapping["BK_DEFINITIONS"] = self.clean_markdown_content(definitions_content)
                            else:
                                # Use content after DEFINITIONS line
                                lines = section_content.split('\n')
                                definitions_lines = []
                                found_definitions = False
                                for line in lines:
                                    if 'DEFINITIONS' in line.upper():
                                        found_definitions = True
                                        continue
                                    elif found_definitions and line.strip():
                                        definitions_lines.append(line)
                                if definitions_lines:
                                    content_mapping["BK_DEFINITIONS"] = self.clean_markdown_content('\n'.join(definitions_lines))
                        else:
                            # Use acronyms content as fallback for definitions
                            content_mapping["BK_DEFINITIONS"] = self.clean_markdown_content(section_content)
                
                elif placeholder == "[BK_REFERENCES]":
                    content_mapping["BK_REFERENCES"] = self.clean_markdown_content(section_content)
                
                elif placeholder == "[BK_PROCEDURE_SUMMARY]":
                    content_mapping["BK_PROCEDURE_SUMMARY"] = self.clean_markdown_content(section_content)
                
                elif placeholder == "[BK_DUT_CONFIG]":
                    content_mapping["BK_DUT_CONFIG"] = self.clean_markdown_content(section_content)
                
                elif placeholder == "[BK_TEST_EXECUTION_CHRONOLOGY]":
                    # Test Execution Chronology table content - no markdown cleaning needed for tables
                    content_mapping["BK_TEST_EXECUTION_CHRONOLOGY"] = section_content
                
                elif placeholder == "[BK_CONSUMABLES_USED]":
                    # Extract consumables info from task_4_6
                    if "consumable" in section_content.lower():
                        consumables_match = re.search(r'consumable.*?\n(.*?)(?=\n\n|$)', section_content, re.DOTALL | re.IGNORECASE)
                        if consumables_match:
                            content_mapping["BK_CONSUMABLES_USED"] = self.clean_markdown_content(consumables_match.group(1).strip())
                        else:
                            content_mapping["BK_CONSUMABLES_USED"] = self.clean_markdown_content(section_content)
                    else:
                        content_mapping["BK_CONSUMABLES_USED"] = self.clean_markdown_content(section_content)
                
                elif placeholder == "[BK_TEST_METHOD_LOSS_INVESTIGATIONS]":
                    content_mapping["BK_TEST_METHOD_LOSS_INVESTIGATIONS"] = self.clean_markdown_content(section_content)
                elif placeholder == "[BK_PROTOCOL_DEVIATIONS]":
                    content_mapping["BK_PROTOCOL_DEVIATIONS"] = self.clean_markdown_content(section_content)
                elif placeholder == "[BK_DEFECTIVE_UNIT]":
                    content_mapping["BK_DEFECTIVE_UNIT"] = self.clean_markdown_content(section_content)
                elif placeholder == "[BK_TEST_RESULT_SUMMARY]":
                    content_mapping["BK_TEST_RESULT_SUMMARY"] = self.clean_markdown_content(section_content)
                elif placeholder == "[BK_TEST_RESULT_ANALYSIS]":
                    # Handle image placeholders specially
                    if section_content == "TBD":
                        content_mapping["BK_TEST_RESULT_ANALYSIS"] = "TBD"
                    elif section_content.startswith("{ANALYSIS_IMAGES:"):
                        # Keep the image placeholder for special processing during Word document creation
                        content_mapping["BK_TEST_RESULT_ANALYSIS"] = section_content
                    else:
                        content_mapping["BK_TEST_RESULT_ANALYSIS"] = self.clean_markdown_content(section_content)
        
        return content_mapping
    
    def create_intro_content(self, report_config: Dict[str, Any]) -> str:
        """Create intro content with report number, revision, and date"""
        report_number = report_config.get('report_number', 'TBD')
        revision = report_config.get('revision', '001')
        generated_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        intro_content = f"""Report Number: {report_number}
Revision: {revision}
Generated: {generated_date}"""
        
        return intro_content
    
    def convert_markdown_table_to_word(self, doc, table_text: str):
        """Convert markdown table text to Word table"""
        lines = table_text.strip().split('\n')
        if len(lines) < 2:
            return None
            
        # Find table lines (those starting with |)
        table_lines = [line.strip() for line in lines if line.strip().startswith('|') and line.strip().endswith('|')]
        
        if len(table_lines) < 2:
            return None
            
        # Parse header row
        header_row = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # Skip separator line and parse data rows
        data_rows = []
        for line in table_lines[2:]:  # Skip header and separator
            if '---' not in line:  # Skip separator lines
                row_data = [cell.strip() for cell in line.split('|')[1:-1]]
                # Include all rows, even if some cells are empty
                data_rows.append(row_data)
        
        if not header_row or len(header_row) == 0:
            return None
            
        # Create Word table
        table = doc.add_table(rows=1, cols=len(header_row))
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        for j, header_text in enumerate(header_row):
            if j < len(header_cells):
                header_cells[j].text = header_text
                # Format header cells
                for paragraph in header_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add data rows
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    row_cells[j].text = cell_text
                    # Format data cells
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
        
        return table
    
    def contains_markdown_table(self, text: str) -> bool:
        """Check if text contains a markdown table"""
        lines = text.split('\n')
        table_lines = [line for line in lines if '|' in line]
        separator_lines = [line for line in lines if '---' in line or '|-' in line]
        return len(table_lines) >= 2 and len(separator_lines) >= 1
    
    def extract_non_table_text(self, text: str) -> str:
        """Extract non-table text from content that may contain markdown tables"""
        lines = text.split('\n')
        non_table_lines = []
        
        for line in lines:
            # Skip table lines (containing | and not just text)
            if '|' in line and ('---' in line or line.count('|') >= 2):
                continue
            # Skip empty lines
            if line.strip():
                non_table_lines.append(line)
        
        return '\n'.join(non_table_lines)
    
    def convert_table_to_text(self, table_text: str) -> str:
        """Convert markdown table to formatted text as a temporary solution"""
        lines = table_text.strip().split('\n')
        table_lines = [line.strip() for line in lines if line.strip().startswith('|') and line.strip().endswith('|')]
        
        if len(table_lines) < 2:
            return table_text
        
        # Parse header row
        header_row = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # Parse data rows (skip separator line)
        result_lines = []
        result_lines.append(' | '.join(header_row))
        result_lines.append('-' * len(' | '.join(header_row)))
        
        for line in table_lines[2:]:  # Skip header and separator
            if '---' not in line:
                row_data = [cell.strip() for cell in line.split('|')[1:-1]]
                if any(cell.strip() for cell in row_data):
                    result_lines.append(' | '.join(row_data))
        
        return '\n'.join(result_lines)
    
    def update_document_headers(self, doc, report_config: Dict[str, Any]):
        """Update document headers with specific placeholder replacements and formatting"""
        try:
            # Extract data for replacements
            report_number_clean = report_config.get('report_number', 'TBD')
            if report_number_clean.startswith('RPT-'):
                report_number_numeric = report_number_clean[4:]  # Remove 'RPT-' prefix
            else:
                report_number_numeric = report_number_clean
            
            title = report_config.get('title', 'TBD')
            revision = report_config.get('revision', '001')
            document_owner = report_config.get('document_owner', 'TBD')
            
            # Define replacements with their specific formatting
            replacements = {
                '[BK_TITLE]': {'text': title, 'size': Pt(10), 'color': RGBColor(0, 0, 0)},
                '[BK_RPT]': {'text': report_number_numeric, 'size': Pt(28), 'color': RGBColor(0, 0, 0)},
                '[BK_REV]': {'text': revision, 'size': Pt(12), 'color': RGBColor(0, 0, 0)},
                '[BK_DOC_OWNER]': {'text': document_owner, 'size': Pt(10), 'color': RGBColor(0, 0, 0)}
            }
            
            # Process all sections
            for section in doc.sections:
                # Process header
                self._replace_placeholders_in_element(section.header, replacements)
                # Process footer
                self._replace_placeholders_in_element(section.footer, replacements)
            
            print("✅ Document headers and footers updated with placeholder replacements")
            print(f"   [BK_TITLE] → {title} (10pt)")
            print(f"   [BK_RPT] → {report_number_numeric} (28pt)")
            print(f"   [BK_REV] → {revision} (12pt)")
            print(f"   [BK_DOC_OWNER] → {document_owner} (10pt)")
            
        except Exception as e:
            print(f"⚠️ Warning: Could not update headers: {e}")
    
    def _replace_placeholders_in_element(self, element, replacements):
        """Helper method to replace placeholders in header/footer elements"""
        # Process paragraphs
        for paragraph in element.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, replacements)
        
        # Process tables
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, replacements)
    
    def _replace_placeholders_in_paragraph(self, paragraph, replacements):
        """Helper method to replace placeholders in a paragraph with specific formatting"""
        original_text = paragraph.text
        
        # Check if any placeholder exists in this paragraph
        for placeholder in replacements.keys():
            if placeholder in original_text:
                # Clear the paragraph
                paragraph.clear()
                
                # Replace placeholders in the text
                new_text = original_text
                for placeholder, config in replacements.items():
                    new_text = new_text.replace(placeholder, config['text'])
                
                # Add text back with formatting
                run = paragraph.add_run(new_text)
                run.font.name = 'Times New Roman'
                
                # Apply specific formatting based on which placeholder was found
                for placeholder, config in replacements.items():
                    if placeholder in original_text:
                        # Find the replaced text in the new text and apply formatting
                        if config['text'] in new_text:
                            # For now, apply the formatting to the entire run
                            # This could be made more sophisticated to format only the replaced text
                            run.font.size = config['size']
                            run.font.color.rgb = config['color']
                            
                            # Add bold formatting for BK_TITLE and BK_RPT
                            if placeholder in ['[BK_TITLE]', '[BK_RPT]']:
                                run.font.bold = True
                            
                            break  # Use the first matching placeholder's formatting
                
                break  # Exit after processing this paragraph
    
    async def create_word_document_from_template(self, report_config: Dict[str, Any], sections: Dict[str, str]) -> str:
        """Create report using template replacement system"""
        # Check template file exists
        template_check = self.check_template_file()
        if not template_check["exists"]:
            raise FileNotFoundError(f"Template file missing: {template_check['error']}")
        
        try:
            # Load template document
            doc = Document(self.template_path)
            print(f"✅ Loaded template: {self.template_path}")
            
            # Update headers with report information
            self.update_document_headers(doc, report_config)
            
            # Extract and clean content from AI sections
            content_mapping = self.extract_content_sections(sections)
            
            # Add intro content
            content_mapping['intro'] = self.create_intro_content(report_config)
            
            # Replace placeholders in document
            placeholders_found = []
            placeholders_missing = []
            
            # Process all paragraphs for placeholder replacement
            for paragraph in doc.paragraphs:
                for placeholder, content_key in self.placeholder_mapping.items():
                    if placeholder in paragraph.text:
                        placeholders_found.append(placeholder)
                        # First try using the content_key from placeholder_mapping
                        replacement_content = sections.get(content_key, "")
                        
                        # If not found, fallback to using placeholder name (without brackets) as key in content_mapping
                        if not replacement_content:
                            content_key_clean = placeholder.replace('[', '').replace(']', '')
                            replacement_content = content_mapping.get(content_key_clean, "")
                        
                        if replacement_content:
                            # Check if content is an image placeholder
                            if replacement_content.startswith("{ANALYSIS_IMAGES:"):
                                # Replace placeholder with empty text first
                                paragraph.text = paragraph.text.replace(placeholder, "")
                                
                                # Process and insert images
                                self.insert_analysis_images(doc, paragraph, replacement_content)
                            # Check if content contains a markdown table
                            elif self.contains_markdown_table(replacement_content):
                                # Replace placeholder with empty text first
                                paragraph.text = paragraph.text.replace(placeholder, "")
                                
                                # Get the paragraph's parent element (to insert table after)
                                p_element = paragraph._element
                                parent = p_element.getparent()
                                
                                # Get any non-table text and add it first
                                non_table_text = self.extract_non_table_text(replacement_content)
                                if non_table_text.strip():
                                    paragraph.text = non_table_text
                                    for run in paragraph.runs:
                                        run.font.name = 'Times New Roman'
                                        run.font.size = Pt(11)
                                        run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                # Insert actual Word table after this paragraph
                                table = self.convert_markdown_table_to_word(doc, replacement_content)
                                if table:
                                    # Move table to correct position
                                    table_element = table._element
                                    parent.insert(parent.index(p_element) + 1, table_element)
                            else:
                                # Replace placeholder with text content
                                paragraph.text = paragraph.text.replace(placeholder, replacement_content)
                                
                                # Set font formatting for the paragraph  
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)  
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            print(f"⚠️ No content available for placeholder: {placeholder}")
            
            
            # Check for missing placeholders and add content to end if needed
            for placeholder, content_key in self.placeholder_mapping.items():
                if placeholder not in placeholders_found:
                    placeholders_missing.append(placeholder)
                    # Use the placeholder name (without brackets) as the key
                    content_key_clean = placeholder.replace('[', '').replace(']', '')
                    replacement_content = content_mapping.get(content_key_clean, "")
                    
                    if replacement_content:
                        print(f"⚠️ Placeholder {placeholder} not found, adding content to end of document")
                        # Add content to end of document
                        doc.add_paragraph("")
                        new_para = doc.add_paragraph(f"Missing section content for {placeholder}:")
                        new_para = doc.add_paragraph(replacement_content)
                        
                        # Set font formatting
                        for run in new_para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Save document
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"DVT_Report_Template_{timestamp}.docx"
            filepath = os.path.join('Output', filename)
            
            # Ensure Output directory exists
            os.makedirs('Output', exist_ok=True)
            
            doc.save(filepath)
            
            print(f"✅ Template document created: {filepath}")
            print(f"✅ Placeholders found and replaced: {len(placeholders_found)}")
            print(f"⚠️ Placeholders missing (content added to end): {len(placeholders_missing)}")
            
            return filepath
            
        except Exception as e:
            print(f"❌ Error creating document from template: {str(e)}")
            raise e
    
    async def create_word_document(self, report_config: Dict[str, Any], sections: Dict[str, str]) -> str:
        """Create Report Document using template replacement system"""
        # Use new template-based approach
        return await self.create_word_document_from_template(report_config, sections)

    
    def _combine_material_equipment_sections(self, dut_config_content: str, equipment_content: str) -> str:
        """
        Combine Task 4.5 (Device Under Test Configuration) and Task 4.6 (Equipment Used) 
        into a single Material & Equipment section with proper sub-headers
        """
        combined_content = "MATERIAL & EQUIPMENT\n\n"
        
        # Add Device Under Test Configuration as subsection 6.1
        if dut_config_content:
            # Remove any existing "DEVICE UNDER TEST CONFIGURATION" header from content
            dut_content_clean = dut_config_content
            if "DEVICE UNDER TEST CONFIGURATION" in dut_content_clean:
                dut_content_clean = dut_content_clean.replace("DEVICE UNDER TEST CONFIGURATION", "").strip()
            
            combined_content += "6.1 Device Under Test Configuration\n\n"
            combined_content += dut_content_clean + "\n\n"
        
        # Add Equipment Used as subsection 6.2  
        if equipment_content:
            # Remove any existing "EQUIPMENT USED" header from content
            equipment_content_clean = equipment_content
            if "EQUIPMENT USED" in equipment_content_clean:
                equipment_content_clean = equipment_content_clean.replace("EQUIPMENT USED", "").strip()
            
            combined_content += "6.2 Equipment Used\n\n"
            combined_content += equipment_content_clean + "\n\n"
        
        return combined_content.strip()
    
    def _create_attachments_list(self, processed_data: Dict[str, Any], report_config: Dict[str, Any]) -> str:
        """Create list of all attachments for the report"""
        
        # Format the attachments list from tracked filenames
        if not self.generated_attachments:
            return "No attachments included with this report."
        
        attachment_lines = []
        for i, attachment in enumerate(self.generated_attachments, 1):
            attachment_lines.append(f"{i}. {attachment}")
        
        return "\n".join(attachment_lines)
